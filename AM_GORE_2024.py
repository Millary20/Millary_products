
"""
*********************************************************************************************
Ayuda Memoria corta con información de presupuestal e información de otras unidades de OPEP

Version Corta

Elaborado: Coordinación de Analistica de Datos y Programación de Presupuesto Territorial
*********************************************************************************************
"""

"""
Lista de los paquetes por instalar:
"""
 
#pip install python-docx ##Paquetes para crear documentos formato "docx"
#pip install nums_from_string ##Paqute para extraer numeros de una cadena
#pip install pyjanitor ##Es una implementación de Python del paquete R janitory proporciona una API limpia para limpiar datos.

import docx
import pandas as pd
import numpy as np
import nums_from_string
import os #Este paquete permite crear actividades dependientes del sistema operativo, por ejemplo crear carpetas, conocer sobre un proceso, finalizar un proceso, etc.
import getpass
import glob
#import matplotlib.pyplot as plt
from datetime import datetime
#from pyprojroot import here
import pyodbc
from janitor import clean_names # pip install pyjanitor
from pathlib import Path
from docx.shared import Pt
from docx.shared import Inches
import re #Nos porporciona opciones de coincidencia.
from docx.shared import Cm # para incluir imagenes en el documento Word
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_UNDERLINE
#import matplotlib.ticker as mtick
from docxtpl import DocxTemplate #pip install docxtpl
from docxtpl import InlineImage
import win32com.client as win32
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

###############################################################################
# Ruta del proyecto #
###############################################################################

#Path: delmódulo pathlib, sirve para manipular rutas locales.
#getuser(): Muestra el nombre de inicio de sesión del usuario. Esta función verifica las variables de entorno LOGNAME, USER, LNAME y USERNAME, en orden, y devuelve el valor de la primera cadena no vacía.
#getpass(): Se utiliza para solicitar a los usuarios que utilicen la solicitud de cadena y lee la entrada del usuario como Contraseña. La lectura de entrada predeterminada es "Contraseña:" y se devuelve a la persona que llama como una cadena

# A continaución se comparan los "Usuarios" de la PC en donde se correra el SCRIPT, se carga la ruta segun el usuario que inicia sesion en la PC

user = getpass.getuser() # Capturo el usuario local / Ejemplo "Llan_"
user.upper()

if   user== "ANALISTAUP29": # PC Analista UP 29 Minedu
     #github = Path(r"C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada")
     proyecto = Path(r"B:\OneDrive - Ministerio de Educación\unidad_B\02_G info\10. Productos\UPP\6. AM automatizada en Python")
     #path_grafico= proyecto + "\graficos"  #Para guardar los graficos 
     #path_mapas= proyecto + "\mapas"  #Para guardar los graficos
     
elif user== "llan_": # PC casa Llan
     #github = Path(r"C:\Users\llan_\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada")
     proyecto = Path(r"C:\Users\llan_\Desktop\AM_AUTOMATIZADA")  
     #path_grafico= proyecto + "\graficos"  #Para guardar los graficos 
     #path_mapas= proyecto + "\mapas"  #Para guardar los graficos  #Para guardar los graficos 
     
else:  # Para cualquier otro usuario
    proyecto = Path(r"C:\Users\Chrystel\Desktop\2024 - MINEDU\MILA_MINEDU_ARCHIVOS\6. AM automatizada en Python")
    
#else:  # Para cualquier otro usuario
    #proyecto = Path(r"C:\Users\analistaup22\Desktop\PRUEBAS\6. AM automatizada en Python")
#else:  # Para cualquier otro usuario
#proyecto = Path(r"B:\OneDrive - Ministerio de Educación\unidad_B\02_G info\10. Productos\UPP\6. AM automatizada en Python")

###############################################################################
# Fechas de corte #
###############################################################################

'''
En esta sección se consigna las fechas que serviran para cargar la información y se crean variables que nos permitan
incorporar información en los textos de la AM.

Cada corte de información actualizada implica que se cambien manualmente los valores de dichas fechas en las variables
que se consignan en esta sección.
'''

# Importamos los nombres de los archivos dentro de la carpeta input
lista_archivos = os.listdir(Path(proyecto, "dataset"))

# Fecha actual
fecha_actual = datetime.today().strftime('%d-%m-%y')
nyear_actual=datetime.today().year
nmeses_actual=datetime.today().month #Rescato el número de mes que estamos hoy
ndia_actual=datetime.today().day
current_quarter = (nmeses_actual - 1) // 3 + 1

## A) Fecha disponibilidad
fecha_corte_disponibilidad = "20240922"
nyear_disponibilidad=fecha_corte_disponibilidad[0:4]
nmes_disponibilidad=fecha_corte_disponibilidad[4:6]
ndia_disponibilidad=fecha_corte_disponibilidad[6:]
fecha_corte_disponibilidad_format=ndia_disponibilidad + "/" + nmes_disponibilidad + "/" + nyear_disponibilidad

nmes_disp_entero=int(nmes_disponibilidad)
nmes_dispo=['enero', 'febrero', 'marzo', 'abril','mayo', 'junio','julio','agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']

## B) Fecha Nexus CAS
fecha_corte_nexus = "20240913"
nyear_nexus=fecha_corte_nexus[0:4]
nmes_nexus=fecha_corte_nexus[4:6]
ndia_nexus=fecha_corte_nexus[6:]
fecha_corte_nexus_2=ndia_nexus + nmes_nexus + nyear_nexus
fecha_corte_nexus_format=ndia_nexus + "/" + nmes_nexus + "/" + nyear_nexus

nmes_nex_entero=int(nmes_nexus)
nmes_nex=['enero', 'febrero', 'marzo', 'abril','mayo', 'junio','julio','agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']

## B) Fecha Nexus PEC
fecha_corte_nexus_3 = "20240913"
nyear_nexus_2=fecha_corte_nexus_3[0:4]
nmes_nexus_2=fecha_corte_nexus_3[4:6]
ndia_nexus_2=fecha_corte_nexus_3[6:]
fecha_corte_nexus_3=ndia_nexus_2 + nmes_nexus_2 + nyear_nexus_2
fecha_corte_nexus_format_2=ndia_nexus_2 + "/" + nmes_nexus_2 + "/" + nyear_nexus_2

nmes_nex_entero_2=int(nmes_nexus_2)
nmes_nex=['enero', 'febrero', 'marzo', 'abril','mayo', 'junio','julio','agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']


'''
## C) Compromisos de desempeño
fecha_corte_compromisos = "20211128"

## D) Fecha remunerativos:
fecha_corte_remune = "20230904"

fecha_corte_remune_formato= datetime.strptime(fecha_corte_remune, "%Y%m%d")
fecha_corte_remune_formato= fecha_corte_remune_formato.strftime("%d.%m.%y")

## D) Fecha mantenimiento:
fecha_corte_mante = "20231012"

nmes_mant=fecha_corte_mante[4:6]
ndia_mant=fecha_corte_mante[6:]
nmes_mant_entero=int(nmes_mant)
nmes_mantenimiento=['enero', 'febrero', 'marzo', 'abril','mayo', 'junio','julio','agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']

fecha_corte_mante= datetime.strptime(fecha_corte_mante, "%Y%m%d")
fecha_corte_mante=fecha_corte_mante.strftime("%Y %m %d")
'''
###############################################################################
# Creación de carpeta donde se guardan los outputs #
###############################################################################

'''
Se crea la carpeta en dondes se almacenará la información correspondiente a las regiones y la fecha de corte.
'''

# Creación de carpeta
# os.path.join() concatena múltiples componentes de una ruta de archivo o directorio
#f"output/CORTA_AM_{fecha_actual}": f"" permite insertar variables dentro de una cadena de texto
#{} permite agregar un valor modificable o dinámico

dir = os.path.join(proyecto, f"output/CORTA_AM_{fecha_actual}")  #Se unen las rutas, f": Lo usamos para concatenar cadenas y valores

if not os.path.exists(dir):
    os.mkdir(dir)  # Para crear una carpeta de acuerdo a lo especificado por "dir"
    print("Se creó una nueva carpeta")
else:
    print("Ya existe la carpeta")

   
# Path de nueva carpeta, crear rutas de directorios
nueva_carpeta = Path(proyecto/ f"output/CORTA_AM_{fecha_actual}")
path_grafico= Path(proyecto/ f"graficos")
path_mapas= Path(proyecto/ f"mapas")

###############################################################################
# Transformación de Datasets                                                  #
###############################################################################

'''
Cargamos la información de las regiones por pliego, para poder capturar esta variable en las demás bases que trabajemos

'''

# Base de datos región
## Cargamos nombres de regiones
nombre_regiones = pd.read_excel(proyecto / "dataset/nombre_regiones.xlsx")

# Base de datos codigo de pliego, ejecutora y ugel
nombre_otros = pd.read_excel(proyecto / "dataset/base_ue_ugel_ubigeo_2023_v2.xlsx")
#nombre_otros.loc[~(nombre_otros['CODOOII'].isnull())]
#nombre_otros.loc[~(((nombre_otros['PLIEGO']==457) & (nombre_otros['EJECUTORA']==301)) | ((nombre_otros['PLIEGO']==464) & (nombre_otros['EJECUTORA']==301)))]
nombre_otros=nombre_otros.drop_duplicates(subset = ['PLIEGO','EJECUTORA'])
nombre_otros = clean_names(nombre_otros) # Normalizamos nombres
nombre_otros.rename(columns={'pliego':'cod_pliego'},inplace=True)
nombre_otros.rename(columns={'ejecutora':'cod_ue'},inplace=True)

###############################
# Base de disponibilidad  #
###############################

'''
La base de disponbilidad contiene información SIAF (PIA, PIM, Devengado, etc) respecto a la ejecución de las intervenciones pedagógicas del MINEDU

El nombre de la base tiene la siguiente estructura:
    
    Disponibilidad_Presupuestal_"yyyy/mm/dd"

1. Se debera de homogenizar las variables de la BD genera y a nivel de intervenciones (las variables deberan tener el mismo nombre)
2. La variable de transferencia para intervenciones en la BD debera tener la siguiente estructura:
  
    tramo1_transferencia ...
    tramo2_transferencia ...
    *
    *
    *
'''

################################
# A) BASE DE DISPONIBILIDAD
################################
## Cargamos base de disponibilidad
data_intervenciones = pd.read_stata(proyecto / f"dataset/bd_disponibilidad/Disponibilidad_Presupuestal_{fecha_corte_disponibilidad}_interv.dta")   
data_intervenciones = clean_names(data_intervenciones) # Normalizamos nombres
data_intervenciones.rename(columns={'intervencion_2':'intervencion'},inplace=True)
data_intervenciones['intervencion'] = data_intervenciones['intervencion'].apply(lambda x: x.strip()) #Para poder quitar los espaicos en blanco antes y despues de las cadenas

# Se quita a PPOR DIT Cunas y SAE
data_intervenciones=data_intervenciones.loc[~((data_intervenciones['cod_intervencion']==53) | (data_intervenciones['cod_intervencion']==54))] 

#FILTRAMOS SOLO INTEREVENCIONES
#*******************************
# Eliminamos filas de "No hay Intervenciones pedagogicas"
data_intervenciones = data_intervenciones[~(data_intervenciones['cod_intervencion']==0)]
# No consideramos: "COAR"
#data_intervenciones = data_intervenciones[~(data_intervenciones['cod_intervencion']==12)]
# Eliminamos  Vacaciones Truncas (Si esque hubiese)
#data_intervenciones = data_intervenciones[data_intervenciones['corr'] != "3.2.8.1.5"]

#------------------------------------------------------------
# TABLA 1: PRESUPUESTO Y EJECUCIÓN A NIVEL DE INTERVENCION
#------------------------------------------------------------  
#data_intervenciones=data_intervenciones.merge(right=nombre_regiones, how="left", on=["cod_pliego"] )
data_intervenciones_t1 = data_intervenciones.groupby(['region','cod_intervencion']).agg({
    'intervencion': 'first',
    'pia_minedu': 'sum',
    'pim': 'sum',
    'devengado': 'sum'
    }).reset_index()

#--------------------------------------------------------------
# TABLA 2: PRESUPUESTO Y EJECUCIÓN A NIVEL DE UNIDAD EJECUTORA
#--------------------------------------------------------------  
#data_intervenciones=data_intervenciones.merge(right=nombre_regiones, how="left", on=["cod_pliego"] )
data_intervenciones_t2 = data_intervenciones.groupby(['region','cod_ue','cod_pliego']).agg({
    'nom_ue': 'first',
    'nom_pliego': 'first',
    'pia_minedu': 'sum',
    'pim': 'sum',
    'devengado': 'sum'
    }).reset_index()

###############################
# Base de nexus PEAS CAS  #
###############################
data_cas_nexus = pd.read_excel(proyecto / f"dataset/bd_plazas_nexus/ReporteNexusCas_IAP_{fecha_corte_nexus_2}.xlsx", sheet_name='Reporte_PEA_2024')   
data_cas_nexus = clean_names(data_cas_nexus) # Normalizamos nombres

data_cas_nexus=data_cas_nexus.loc[~((data_cas_nexus['cod_int']==53) | (data_cas_nexus['cod_int']==54))] 
#data_cas_nexus.rename(columns={'intervencion_nombre_corto':'intervencion_fake'},inplace=True)
#data_cas_nexus.rename(columns={'intervencion':'intervencion_nombre_corto'},inplace=True)

#----------------------------------------------------------------------------
# TABLA 3: PEA VACANTE, CONTRATADA Y PROGRAMADA A NIVEL DE UNIDAD EJECUTORA
#----------------------------------------------------------------------------  

data_cas_nexus_t1 = data_cas_nexus.groupby(['region','cod_ue']).agg({
    'unidad_ejecutora': 'first',
    'contratado': 'sum',
    'vacante': 'sum',
    'programadas': 'sum'
    }).reset_index()

###############################
# Base de nexus PEC           #
###############################
data_pec_nexus = pd.read_excel(proyecto / f"dataset/bd_plazas_nexus/ReporteNexus_PEC_PRONOEI_{fecha_corte_nexus_3}.xlsx", sheet_name='Reporte_PEA_2024')   
data_pec_nexus = clean_names(data_pec_nexus) # Normalizamos nombres
data_pec_nexus=data_pec_nexus.loc[~((data_pec_nexus['cod_int']==53) | (data_pec_nexus['cod_int']==54))] 

data_pec_nexus['ruralidad'].unique()

#data_cas_nexus.rename(columns={'intervencion_nombre_corto':'intervencion_fake'},inplace=True)
#data_cas_nexus.rename(columns={'intervencion':'intervencion_nombre_corto'},inplace=True)

data_pec_nexus['rural']=0
data_pec_nexus.loc[data_pec_nexus['ruralidad']=='RURAL', 'rural' ]=1

data_pec_nexus['urbano']=0
data_pec_nexus.loc[data_pec_nexus['ruralidad']=='URBANO', 'urbano' ]=1

#----------------------------------------------------------------------------
# TABLA 4: PEC RURALIDAD
#----------------------------------------------------------------------------  

data_cas_nexus_t2 = data_pec_nexus.groupby(['region','cod_ue']).agg({
    'unidad_ejecutora': 'first',
    'contratado': 'sum',
    'vacante': 'sum',
    'programadas': 'sum',
    'rural' : 'sum',
    'urbano': 'sum'
    }).reset_index()

'''
##################################
# Base de CDD                    #
##################################
data_cdd = pd.read_excel(proyecto / "dataset/bd_cdd/Anexo C.xlsx", sheet_name='Anexo C', skiprows=7,  nrows=226, usecols = 'B:E') #,skiprows=0,  nrows=249, usecols = 'B:j') 
data_cdd = clean_names(data_cdd) # Normalizamos nombres
data_cdd=data_cdd.fillna(0)

#data_punche=data_punche.rename(columns={"departamento": "region"}) 
data_cdd=data_cdd[['region','unidad_ejecutora_de_educacion', 'transferencia_por_cumplimiento']].groupby(by=['region','unidad_ejecutora_de_educacion'], as_index=False).sum()
'''
'''
# Funcion para crear links dentro del word:
def add_hyperlink(paragraph, url, text):
    """
    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Add blue color to the hyperlink
    color_element = docx.oxml.shared.OxmlElement('w:color')
    color_element.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(color_element)

    # Add underline to the hyperlink
    u_element = docx.oxml.shared.OxmlElement('w:u')
    u_element.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u_element)
    
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    
    return hyperlink
'''
###############################################################################
# Creación del documento en docx 
###############################################################################

# Generamos la lista de Regiones
#lista_regiones = ["AMAZONAS", "ANCASH", "APURIMAC", "AREQUIPA", "AYACUCHO", "CAJAMARCA", "CUSCO", "HUANCAVELICA", "HUANUCO", "ICA", "JUNIN", "LA LIBERTAD", "LAMBAYEQUE", "LORETO", "MADRE DE DIOS", "MOQUEGUA", "PASCO", "PIURA", "PUNO", "SAN MARTIN", "TACNA", "TUMBES", "UCAYALI", "LIMA PROVINCIAS", "CALLAO", "LIMA METROPOLITANA"]
#lista_regiones = ["AMAZONAS", "ANCASH"]
# For loop para cada región
#for region in lista_regiones:
region = "AMAZONAS"
###############################################################################
# 1. Construcción de tablas e indicadores                                      #
###############################################################################
   
    #Año actual: nyear_actual
year_actual = str(nyear_actual)
    
    ##################################################################################
    # TABLA 1: PRESUPUESTO Y EJECUCIÓN A NIVEL DE INTERVENCION
    ##################################################################################
region_seleccionada = data_intervenciones_t1['region'] == region #Seleccionar region
tabla_intervenciones = data_intervenciones_t1[region_seleccionada]
    
    #INDICADORES PARA EL TEXTO
    #Cuenta cuantas intervenciones tiene la region
cant_iap = str('{:,.0f}'.format(tabla_intervenciones["cod_intervencion"].count()))
    
tabla_intervenciones_formato_1=tabla_intervenciones.copy()
  
    # Generamos porcentaje de avance
tabla_intervenciones_formato_1['ejecucion'] = tabla_intervenciones_formato_1["devengado"]/tabla_intervenciones_formato_1["pim"]
tabla_intervenciones_formato_1.loc[tabla_intervenciones_formato_1['ejecucion'].isnull(),'ejecucion']=0
tabla_intervenciones_formato_1=tabla_intervenciones_formato_1.sort_values('ejecucion', ascending=False)

    # Generamos fila total
total_ejec = tabla_intervenciones_formato_1["devengado"].sum()/tabla_intervenciones_formato_1["pim"].sum() if (tabla_intervenciones_formato_1["devengado"].sum()/tabla_intervenciones_formato_1["pim"].sum())>0 else 0
totalint = tabla_intervenciones_formato_1[['pia_minedu','pim','devengado','ejecucion','region']].groupby(by = ["region"], as_index=False).sum()

    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_1 = tabla_intervenciones_formato_1.append(totalint, ignore_index=True)
tabla_intervenciones_formato_1 = pd.concat([tabla_intervenciones_formato_1, totalint])
tabla_intervenciones_formato_1.reset_index(drop=True, inplace=True)
tabla_intervenciones_formato_1['intervencion'] = tabla_intervenciones_formato_1['intervencion'].fillna("Total")
tabla_intervenciones_formato_1['ejecucion'] = tabla_intervenciones_formato_1['ejecucion'].replace(np.inf, 0)
tabla_intervenciones_formato_1['ejecucion'] = tabla_intervenciones_formato_1['ejecucion'].fillna("0").astype(float)
# iloc[] permite seleccionar filas y columnas por índices, el -1 significa comenzar de la última fila
tabla_intervenciones_formato_1.iloc[-1, tabla_intervenciones_formato_1.columns.get_loc('ejecucion')] = total_ejec
    
    # Formato para la tabla de PIM y Devengado por intervención
tabla_intervenciones_formato_1=tabla_intervenciones_formato_1[['intervencion','pia_minedu','pim','devengado','ejecucion']]
    
    # Formato para la tabla
formato_tabla_intervenciones = {
    "intervencion" : "{}",
    "pia_minedu" : "{:,.0f}",
    "pim" : "{:,.0f}",
    "devengado": "{:,.0f}",
    "ejecucion": "{:,.1%}"
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
tabla_intervenciones_formato_1 = tabla_intervenciones_formato_1.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
    
    ##################################################################################
    # TABLA 2: PRESUPUESTO Y EJECUCIÓN A NIVEL DE UNIDAD EJECUTORA
    ##################################################################################
region_seleccionada = data_intervenciones_t2['region'] == region #Seleccionar region
tabla_intervenciones = data_intervenciones_t2[region_seleccionada]
    
    #INDICADORES PARA EL TEXTO
    #info_pliego = tabla_intervenciones['nom_pliego'][0]
valores_unicos = tabla_intervenciones['nom_pliego'].unique()
nombre_pliego = str(valores_unicos[0])
nombre_pliego = nombre_pliego.strip("['']")
    
tabla_intervenciones_formato_2=tabla_intervenciones.copy()
    
    # Generamos porcentaje de avance
tabla_intervenciones_formato_2['ejecucion'] = tabla_intervenciones_formato_2["devengado"]/tabla_intervenciones_formato_2["pim"]
tabla_intervenciones_formato_2.loc[tabla_intervenciones_formato_2['ejecucion'].isnull(),'ejecucion']=0
tabla_intervenciones_formato_2=tabla_intervenciones_formato_2.sort_values('ejecucion', ascending=False)

    # Generamos fila total
total_ejec = tabla_intervenciones_formato_2["devengado"].sum()/tabla_intervenciones_formato_2["pim"].sum() if (tabla_intervenciones_formato_2["devengado"].sum()/tabla_intervenciones_formato_2["pim"].sum())>0 else 0
totalint = tabla_intervenciones_formato_2[['pia_minedu','pim','devengado','ejecucion','region']].groupby(by = ["region"], as_index=False).sum()

    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_1 = tabla_intervenciones_formato_1.append(totalint, ignore_index=True)
tabla_intervenciones_formato_2 = pd.concat([tabla_intervenciones_formato_2, totalint])
tabla_intervenciones_formato_2.reset_index(drop=True, inplace=True)
tabla_intervenciones_formato_2['nom_ue'] = tabla_intervenciones_formato_2['nom_ue'].fillna("Total")
tabla_intervenciones_formato_2['ejecucion'] = tabla_intervenciones_formato_2['ejecucion'].replace(np.inf, 0)
tabla_intervenciones_formato_2['ejecucion'] = tabla_intervenciones_formato_2['ejecucion'].fillna("0").astype(float)
tabla_intervenciones_formato_2.iloc[-1, tabla_intervenciones_formato_2.columns.get_loc('ejecucion')] = total_ejec
    
    # Formato para la tabla de PIM y Devengado por intervención
tabla_intervenciones_formato_2=tabla_intervenciones_formato_2[['nom_ue','pia_minedu','pim','devengado','ejecucion']]
    
    # Formato para la tabla
formato_tabla_intervenciones = {
    "nom_ue" : "{}",
    "pia_minedu" : "{:,.0f}",
    "pim" : "{:,.0f}",
    "devengado": "{:,.0f}",
    "ejecucion": "{:,.1%}"
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
tabla_intervenciones_formato_2 = tabla_intervenciones_formato_2.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
    
    ##################################################################################
    # TABLA 3: AVANCE DE CONTRATACIÓN DE PLAZAS NEXUS
    ##################################################################################
region_seleccionada = data_cas_nexus_t1['region'] == region #Seleccionar region
tabla_intervenciones = data_cas_nexus_t1[region_seleccionada]
  
    #INDICADORES PARA EL TEXTO
ejecucion = str('{:,.1%}'.format(tabla_intervenciones["contratado"].sum()/tabla_intervenciones["programadas"].sum())  if tabla_intervenciones["contratado"].sum()/tabla_intervenciones["programadas"].sum()>0 else 0)
if ejecucion=="0":
        ejecucion ="0.0%"
    
tabla_intervenciones_formato_3=tabla_intervenciones.copy()
    
    # Generamos porcentaje de avance
tabla_intervenciones_formato_3['avace'] = tabla_intervenciones_formato_3["contratado"]/tabla_intervenciones_formato_3["programadas"]
tabla_intervenciones_formato_3.loc[tabla_intervenciones_formato_3['avace'].isnull(),'avace']=0
tabla_intervenciones_formato_3=tabla_intervenciones_formato_3.sort_values('avace', ascending=False)

    # Generamos fila total
total_ejec = tabla_intervenciones_formato_3["contratado"].sum()/tabla_intervenciones_formato_3["programadas"].sum() if (tabla_intervenciones_formato_3["contratado"].sum()/tabla_intervenciones_formato_3["programadas"].sum())>0 else 0
totalint = tabla_intervenciones_formato_3[['contratado','vacante','programadas','avace','region']].groupby(by = ["region"], as_index=False).sum()

    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_1 = tabla_intervenciones_formato_1.append(totalint, ignore_index=True)
tabla_intervenciones_formato_3 = pd.concat([tabla_intervenciones_formato_3, totalint])
tabla_intervenciones_formato_3.reset_index(drop=True, inplace=True)
tabla_intervenciones_formato_3['unidad_ejecutora'] = tabla_intervenciones_formato_3['unidad_ejecutora'].fillna("Total")
tabla_intervenciones_formato_3['avace'] = tabla_intervenciones_formato_3['avace'].replace(np.inf, 0)
tabla_intervenciones_formato_3['avace'] = tabla_intervenciones_formato_3['avace'].fillna("0").astype(float)
tabla_intervenciones_formato_3.iloc[-1, tabla_intervenciones_formato_3.columns.get_loc('avace')] = total_ejec
    
    # Formato para la tabla de PIM y Devengado por intervención
tabla_intervenciones_formato_3=tabla_intervenciones_formato_3[['unidad_ejecutora','contratado','vacante','programadas','avace']]
    
    # Formato para la tabla
formato_tabla_intervenciones = {
    "unidad_ejecutora" : "{}",
    "contratado" : "{:,.0f}",
    "vacante" : "{:,.0f}",
    "programadas": "{:,.0f}",
    "avace": "{:,.1%}"
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
tabla_intervenciones_formato_3 = tabla_intervenciones_formato_3.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
   
    ##################################################################################
    # TABLA 4: AVANCE DE CONTRATACIÓN DE PEC NEXUS
    ##################################################################################
region_seleccionada = data_cas_nexus_t2['region'] == region #Seleccionar region
tabla_intervenciones = data_cas_nexus_t2[region_seleccionada]
    
    #INDICADORES PARA EL TEXTO
ejecucion_2 = str('{:,.1%}'.format(tabla_intervenciones["contratado"].sum()/tabla_intervenciones["programadas"].sum())  if tabla_intervenciones["contratado"].sum()/tabla_intervenciones["programadas"].sum()>0 else 0)
if ejecucion_2=="0":
   ejecucion_2 ="0.0%"
    
tabla_intervenciones_formato_4=tabla_intervenciones.copy()
    
    # Generamos porcentaje de avance
tabla_intervenciones_formato_4['avace'] = tabla_intervenciones_formato_4["contratado"]/tabla_intervenciones_formato_4["programadas"]
tabla_intervenciones_formato_4.loc[tabla_intervenciones_formato_4['avace'].isnull(),'avace']=0
tabla_intervenciones_formato_4=tabla_intervenciones_formato_4.sort_values('avace', ascending=False)

    # Generamos fila total
total_ejec = tabla_intervenciones_formato_4["contratado"].sum()/tabla_intervenciones_formato_4["programadas"].sum() if (tabla_intervenciones_formato_4["contratado"].sum()/tabla_intervenciones_formato_4["programadas"].sum())>0 else 0
totalint = tabla_intervenciones_formato_4[['contratado','vacante','programadas','avace','rural','urbano','region']].groupby(by = ["region"], as_index=False).sum()

    # Realizamos append del total en la tabla
    #tabla_intervenciones_formato_1 = tabla_intervenciones_formato_1.append(totalint, ignore_index=True)
tabla_intervenciones_formato_4 = pd.concat([tabla_intervenciones_formato_4, totalint])
tabla_intervenciones_formato_4.reset_index(drop=True, inplace=True)
tabla_intervenciones_formato_4['unidad_ejecutora'] = tabla_intervenciones_formato_4['unidad_ejecutora'].fillna("Total")
tabla_intervenciones_formato_4['avace'] = tabla_intervenciones_formato_4['avace'].replace(np.inf, 0)
tabla_intervenciones_formato_4['avace'] = tabla_intervenciones_formato_4['avace'].fillna("0").astype(float)
tabla_intervenciones_formato_4.iloc[-1, tabla_intervenciones_formato_4.columns.get_loc('avace')] = total_ejec
    
    # Formato para la tabla de PIM y Devengado por intervención
tabla_intervenciones_formato_4=tabla_intervenciones_formato_4[['unidad_ejecutora','contratado','vacante','programadas','avace','rural','urbano']]
    
    # Formato para la tabla
formato_tabla_intervenciones = {
    "unidad_ejecutora" : "{}",
    "programadas": "{:,.0f}",
    "rural": "{:,.0f}",
    "urbano": "{:,.0f}",
    "contratado" : "{:,.0f}",
    "vacante" : "{:,.0f}",
    "avace": "{:,.1%}"
    #    "costo_actual": "{:,.0f}",
    #    "ejecucion": "{:,.1%}",
    }
    
tabla_intervenciones_formato_4 = tabla_intervenciones_formato_4.transform({k: v.format for k, v in formato_tabla_intervenciones.items()})
    
###############################################################################
# 3. Inclusión del texto del documento                                        #
###############################################################################
region_titulo=region
region=region.lower().capitalize() 
#document = Document(proyecto / "formato/FORMATO_FINAL_UPP_CORTA.docx") # Creación del documento en base al template
document = Document(proyecto / "formato/FORMATO_FINAL_UPP_CORTA - copia.docx") # Creación del documento en base al template   
    #####################################################
    # TITULO GENERAL               
    ####################################################
titulo_general =document.add_heading(f"AYUDA MEMORIA\nREGIÓN {region_titulo}", level=1)
titulo_paragraph_format = titulo_general.paragraph_format
titulo_paragraph_format.space_before = Pt(0)
titulo_paragraph_format.space_after = Pt(0) 
titulo_general.paragraph_format.space_before = Pt(0)


    #####################################################
    # 1. INTERVENCIONES Y ACCIONES PEDAGÓGICAS         
    ####################################################

# Agregar un encabezado de nivel 2
heading = document.add_heading("INTERVENCIONES Y ACCIONES PEDAGÓGICAS", level=2)
# Alinear el párrafo al centro
heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Ajustar el formato del párrafo para ese encabezado
heading.paragraph_format.space_before = Pt(12)  # Espacio antes del encabezado

region_1 = region.capitalize()

interv_parrafo2 = document.add_paragraph(f"Los créditos presupuestarios asignados para el financiamiento de las intervenciones y acciones pedagógicas a implementarse en la región {region_1}\
 para el presente ejercicio fiscal, se dan de conformidad con los numerales 94.1, 94.2 y 94.3 del artículo 94 de la Ley N° 31953, Ley de Presupuesto del Sector Público para el Año Fiscal 2024,\
 asimismo, cabe señalar que se está incluyendo también el financiamiento de la intervención de PPOR DIT – EduCunas, en el marco del artículo 33 de la Ley N° 31953.")
interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY     

document.add_heading("I. NIVEL DE EJECUCIÓN PRESUPUESTAL", level=3) # 1) Intervenciones pedagógicas  
    
interv_parrafo = document.add_paragraph(f"A nivel de intervención y acción pedagógica, se precisa que la región {region_1} implementaría este año fiscal {year_actual}, un total de {cant_iap} intervenciones, según\
 se enlista a continuación:")
#interv_parrafo = document.add_paragraph(f"A nivel de intervención y acción pedagógica, se precisa que la región {region_1} implementaría este año fiscal {year_actual}, un total de {cant_iap} intervenciones, según\
#se enlista a continuación:", style='List Bullet')
interv_parrafo.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY      
interv_parrafo.paragraph_format.space_after = Pt(12)

# Crear el párrafo y agregar el texto
titulo_tabla1 = document.add_paragraph()

# Agregar un run para el texto del título
run = titulo_tabla1.add_run("Tabla 1: Nivel de avance de ejecución presupuestal por intervención")

# Hacer el texto en negrita
run.bold = True

# Alinear el párrafo al centro
titulo_tabla1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Establecer el espacio antes del párrafo
#titulo_tabla1.paragraph_format.space_before = Pt(1)

    ###########################################################################################################################
    # TABLA 1: PRESUPUESTO Y EJECUCIÓN A NIVEL DE INTERVENCION
    ###########################################################################################################################   
tabla1_interv = document.add_table(tabla_intervenciones_formato_1.shape[0]+1, tabla_intervenciones_formato_1.shape[1])
tabla1_interv.autofit = False
tabla1_interv.allow_autofit = True
tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
row = tabla1_interv.rows[0].cells
row[0].text = "Intervención"
row[1].text = "PIA Asignado"
row[2].text = "PIM SIAF"
row[3].text = "Devengado"
row[4].text = "Avance de ejecución"
    #row[3].text = "DEV."
    #row[4].text = "% DEV"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
## Contenido de la tabla
for i in range(tabla_intervenciones_formato_1.shape[0]):
    for j in range(tabla_intervenciones_formato_1.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_1.values[i,j])
            
font_size = Pt(10)
for row in tabla1_interv.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = font_size  # Establecer el tamaño de fuente
                run.font.name = 'Arial Unicode MS'  # Cambiar la fuente a Arial            
   
for row in tabla1_interv.rows[1:]:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT #columna 1
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 2
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 3
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 4
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 5
    
for celda in tabla1_interv.rows[-1].cells:
        for paragraph in celda.paragraphs:
            paragraph.alignment = 1  # 0: Izquierda, 1: Centrado, 2: Derecha, 3: Justificado
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(89, 89, 89)

for row_idx, row in enumerate(tabla1_interv.rows):
    # Aquí puedes hacer referencia a row_idx, que es un entero
    tabla1_interv.cell(row_idx, 0).width = Inches(2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 1).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 2).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 3).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 4).width = Inches(1.2)  # Establecer el ancho solo para la columna 0  
   
    
    #Mi nuevo estilo para fuente:
fuente_stilo = document.styles.add_style('Fuentes', WD_STYLE_TYPE.PARAGRAPH)
fuente_stilo.font.name = 'Arial (Cuerpo)'
fuente_stilo.font.size = docx.shared.Pt(9)

    #Fuente:
interv_parrafo = document.add_paragraph(f"Fuente: Base SIAF al {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} de {year_actual}.")
interv_parrafo.style = fuente_stilo
interv_parrafo.style.font.italic = True
interv_parrafo.style.font.size = docx.shared.Pt(9)      
interv_parrafo.paragraph_format.space_after = Pt(12)
    
    #document.add_paragraph('')
    
    #####################################################
    # SEGUNDO PARRAFO
    ####################################################
region_1 = region.capitalize()
           
interv_parrafo3 = document.add_paragraph(f"De igual manera, en el siguiente cuadro se muestra la información del marco presupuestal asignado distribuido por unidad ejecutora\
 del Pliego {nombre_pliego}:")
interv_parrafo3.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY  
interv_parrafo3.paragraph_format.space_before = Pt(12)
interv_parrafo3.paragraph_format.space_after = Pt(12)

# Crear el párrafo y agregar el texto
titulo_tabla2 = document.add_paragraph()

# Agregar un run para el texto del título
run2 = titulo_tabla2.add_run("Tabla 2: Nivel de avance de ejecución presupuestal por Región")

# Hacer el texto en negrita
run2.bold = True

# Alinear el párrafo al centro
titulo_tabla2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ###########################################################################################################################
    # TABLA 2: PRESUPUESTO Y EJECUCIÓN A NIVEL DE UNIDAD EJECUTORA
    ###########################################################################################################################     
tabla1_interv = document.add_table(tabla_intervenciones_formato_2.shape[0]+1, tabla_intervenciones_formato_2.shape[1])
tabla1_interv.autofit = False
tabla1_interv.allow_autofit = True
tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
row = tabla1_interv.rows[0].cells
row[0].text = "Unidad Ejecutora"
row[1].text = "PIA asignado"
row[2].text = "PIM SIAF"
row[3].text = "Devengado"
row[4].text = "Avance de Ejecución"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
for i in range(tabla_intervenciones_formato_2.shape[0]):
        for j in range(tabla_intervenciones_formato_2.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_2.values[i,j])
                
font_size = Pt(10)
for row in tabla1_interv.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = font_size  # Establecer el tamaño de fuente
                run.font.name = 'Arial Unicode MS'  # Cambiar la fuente a Arial       
   
    
for row in tabla1_interv.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT #columna 1
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 2
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 3
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 4
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 5
    
for celda in tabla1_interv.rows[-1].cells:
        for paragraph in celda.paragraphs:
            paragraph.alignment = 1  # 0: Izquierda, 1: Centrado, 2: Derecha, 3: Justificado
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(89, 89, 89)         

for row_idx, row in enumerate(tabla1_interv.rows):
    # Aquí puedes hacer referencia a row_idx, que es un entero
    tabla1_interv.cell(row_idx, 0).width = Inches(2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 1).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 2).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 3).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 4).width = Inches(1.2)  # Establecer el ancho solo para la columna 0              

'''
    #Mi nuevo estilo para fuente:
    fuente_stilo = document.styles.add_style('Fuentes', WD_STYLE_TYPE.PARAGRAPH)
    fuente_stilo.font.name = 'Arial (Cuerpo)'
    fuente_stilo.font.size = docx.shared.Pt(9)
'''
    #Fuente:
interv_parrafo = document.add_paragraph(f"Fuente: Base SIAF al {ndia_disponibilidad} de {nmes_dispo[nmes_disp_entero-1]} de {year_actual}.")
interv_parrafo.style = fuente_stilo
interv_parrafo.style.font.italic = True
interv_parrafo.style.font.size = docx.shared.Pt(9)     
interv_parrafo.paragraph_format.space_after = Pt(12)    
    #document.add_paragraph('')
    
    #####################################################
    # INFORMACIÓN DE PLAZAS CAS NEXUS
    ####################################################
    
document.add_heading("II. NIVEL DE AVANCE DE CONTRATACIÓN", level=3) # 1) Intervenciones pedagógicas 

# Crear el encabezado
heading = document.add_heading("2.1. Nivel de avance de contratación de plazas CAS", level=3)

# Acceder al run del encabezado
run = heading.runs[0]

# Cambiar el color de la fuente (por ejemplo, a rojo)
run.font.color.rgb = RGBColor(172, 108, 27)  # Rojo (RGB)

region_1 = region.capitalize()
interv_parrafo2 = document.add_paragraph(f"La región {region_1} presenta un nivel de avance de contratación del {ejecucion} de las plazas\
 programadas en el marco de las intervenciones pedagógicas, cuya distribución se muestra según el siguiente resumen:")
interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY     
interv_parrafo2.paragraph_format.space_after = Pt(12)   
    
    ###########################################################################################################################
    # TABLA 3: AVANCE DE CONTRATACIÓN DE PLAZAS NEXUS
    ###########################################################################################################################     

# Crear el párrafo y agregar el texto
titulo_tabla3 = document.add_paragraph()
titulo_tabla3.paragraph_format.space_before = Pt(12)   
# Agregar un run para el texto del título
run3 = titulo_tabla3.add_run("Tabla 3: Nivel de avance de contratación de plazas CAS según Nexus")

# Hacer el texto en negrita
run3.bold = True

# Alinear el párrafo al centro
titulo_tabla3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

tabla1_interv = document.add_table(tabla_intervenciones_formato_3.shape[0]+1, tabla_intervenciones_formato_3.shape[1])
tabla1_interv.autofit = False
tabla1_interv.allow_autofit = True
tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
row = tabla1_interv.rows[0].cells
row[0].text = "Unidad Ejecutora"
row[1].text = "PEA Contratada"
row[2].text = "PEA Vacante"
row[3].text = "PEA Programada"
row[4].text = "% Avance Contratación"
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
for i in range(tabla_intervenciones_formato_3.shape[0]):
        for j in range(tabla_intervenciones_formato_3.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_3.values[i,j])
             
font_size = Pt(10)
for row in tabla1_interv.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = font_size  # Establecer el tamaño de fuente
                run.font.name = 'Arial Unicode MS'  # Cambiar la fuente a Arial     
   
for row in tabla1_interv.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT #columna 1
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 2
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 3
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 4
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 5
    
for celda in tabla1_interv.rows[-1].cells:
        for paragraph in celda.paragraphs:
            paragraph.alignment = 1  # 0: Izquierda, 1: Centrado, 2: Derecha, 3: Justificado
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(89, 89, 89)         
for row_idx, row in enumerate(tabla1_interv.rows):
    # Aquí puedes hacer referencia a row_idx, que es un entero
    tabla1_interv.cell(row_idx, 0).width = Inches(2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 1).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 2).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 3).width = Inches(1.2)  # Establecer el ancho solo para la columna 0
    tabla1_interv.cell(row_idx, 4).width = Inches(1.2)  # Establecer el ancho solo para la columna 0           
'''
    #Mi nuevo estilo para fuente:
    fuente_stilo = document.styles.add_style('Fuentes', WD_STYLE_TYPE.PARAGRAPH)
    fuente_stilo.font.name = 'Arial (Cuerpo)'
    fuente_stilo.font.size = docx.shared.Pt(9)
'''
    #Fuente:
interv_parrafo = document.add_paragraph(f"Fuente: Sistema Nexus al {ndia_nexus} de {nmes_nex[nmes_nex_entero-1]} de {year_actual}.")
interv_parrafo.style = fuente_stilo
interv_parrafo.style.font.italic = True
interv_parrafo.style.font.size = docx.shared.Pt(9) 
interv_parrafo.paragraph_format.space_after = Pt(12)       
    #####################################################
    # INFORMACIÓN DE PEC NEXUS
    ####################################################

# Crear el encabezado
heading = document.add_heading("2.2. Nivel de avance de contratación de PEC", level=3)

# Acceder al run del encabezado
run = heading.runs[0]

# Cambiar el color de la fuente (por ejemplo, a rojo)
run.font.color.rgb = RGBColor(172, 108, 27)  # Rojo (RGB)

  
region_1 = region.capitalize()
interv_parrafo2 = document.add_paragraph(f"La región {region_1} presenta un nivel de avance de contratación del {ejecucion_2} de las PEC\
 programadas en el marco de las intervenciones pedagógicas denominada PRONOEI, cuya distribución se muestra según el siguiente resumen:")
interv_parrafo2.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY       
interv_parrafo2.paragraph_format.space_after = Pt(12)       
    ###########################################################################################################################
    # TABLA 4: AVANCE DE CONTRATACIÓN DE PLAZAS NEXUS
    ###########################################################################################################################     

# Crear el párrafo y agregar el texto
titulo_tabla4 = document.add_paragraph()
titulo_tabla4.paragraph_format.space_before = Pt(12)   
# Agregar un run para el texto del título
run4 = titulo_tabla4.add_run("Tabla 4: Nivel de avance de contratación de PEC según Nexus")

# Hacer el texto en negrita
run4.bold = True

# Alinear el párrafo al centro
titulo_tabla4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

tabla1_interv = document.add_table(tabla_intervenciones_formato_4.shape[0]+1, tabla_intervenciones_formato_4.shape[1])
tabla1_interv.autofit = False
tabla1_interv.allow_autofit = True
tabla1_interv.style = "tabla_minedu_1"
    #tabla1_interv.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
tabla1_interv.alignment = WD_TABLE_ALIGNMENT.CENTER
    
row = tabla1_interv.rows[0].cells
row[0].text = "Unidad Ejecutora"
row[1].text = "PEC Programada"
row[2].text = "Rural"
row[3].text = "Urbano"
row[4].text = "PEC Contratada"
row[5].text = "PEC Vacante"
row[6].text = "% Avance Contratación"
    
    #row[5].text = "COSTO AL MES"
    #row[4].text = "% DEV COSTO AL MES"
    ## Contenido de la tabla
for i in range(tabla_intervenciones_formato_4.shape[0]):
        for j in range(tabla_intervenciones_formato_4.shape[-1]):
            tabla1_interv.cell(i+1,j).text = str(tabla_intervenciones_formato_4.values[i,j])
             
font_size = Pt(10)
for row in tabla1_interv.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = font_size  # Establecer el tamaño de fuente
                run.font.name = 'Arial Unicode MS'  # Cambiar la fuente a Arial 
   
for row in tabla1_interv.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT #columna 1
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 2
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 3
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 4
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 5
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 6
        row.cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT #Columan 7
    
for celda in tabla1_interv.rows[-1].cells:
        for paragraph in celda.paragraphs:
            paragraph.alignment = 1  # 0: Izquierda, 1: Centrado, 2: Derecha, 3: Justificado
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(89, 89, 89)         

tabla1_interv.cell(row_idx, 0).width = Inches(1)  # Establecer el ancho solo para la columna 0
tabla1_interv.cell(row_idx, 1).width = Inches(0.5)  # Establecer el ancho solo para la columna 0
tabla1_interv.cell(row_idx, 2).width = Inches(0.5)  # Establecer el ancho solo para la columna 0
tabla1_interv.cell(row_idx, 3).width = Inches(0.5)  # Establecer el ancho solo para la columna 0
tabla1_interv.cell(row_idx, 4).width = Inches(0.5)  # Establecer el ancho solo para la columna 0 
tabla1_interv.cell(row_idx, 5).width = Inches(0.5)  # Establecer el ancho solo para la columna 0 
tabla1_interv.cell(row_idx, 6).width = Inches(0.5)  # Establecer el ancho solo para la columna 0             
'''
    #Mi nuevo estilo para fuente:
    fuente_stilo = document.styles.add_style('Fuentes', WD_STYLE_TYPE.PARAGRAPH)
    fuente_stilo.font.name = 'Arial (Cuerpo)'
    fuente_stilo.font.size = docx.shared.Pt(9)
'''
    #Fuente:
interv_parrafo = document.add_paragraph(f"Fuente: Sistema Nexus al {ndia_nexus_2} de {nmes_nex[nmes_nex_entero_2-1]} de {year_actual}.")
interv_parrafo.style = fuente_stilo
interv_parrafo.style.font.italic = True
interv_parrafo.style.font.size = docx.shared.Pt(9) 
    
    #######################
    # Guardamos documento #
    #######################
document.save(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
'''
    #Inseraremos inforamcion de algunas variables que tenemos en word
    #--------------------------------------------------------------------
    #El nombre de region en la caratural, los resumenes de variables y el gráfico de mapa
    doc = DocxTemplate(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
    path_mapa=path_mapas / f"{region}.png"
    path_mapa = path_mapa.as_posix()

    mi_region=region
    mi_ejecutora=cant_ue_region
    cost_interven=pia_intervenciones_region
    fecha_siaf=fecha_corte_disponibilidad_format
    fecha_nexus=fecha_corte_nexus_format
    transf_materia=transf_materiales
    transf_remunerativa=0 #Por ahora estoy considerando 0 soles
    
    context = {'mi_region': mi_region, 'mi_ejecutora': mi_ejecutora, 'cost_interven':cost_interven,
               'fecha_siaf':fecha_siaf, 'fecha_nexus':fecha_nexus, 'mi_imagen':InlineImage(doc, path_mapa, width=Cm(7.38), height=Cm(9.82)),
               'transf_materia': transf_materia, 'tranf_remunerativas':transf_remunerativa}
    
    doc.render(context)
    doc.save(nueva_carpeta / f'AM_{region}_{fecha_actual}.docx')
    
    #Finalmente cargaremos el word para actualizar el indice 
    #--------------------------------------------------------------------
    
    # Crear una instancia de la aplicación Word
    word = win32.Dispatch('Word.Application')

    # Abrir el archivo de Word y no activamos su ventana
    file_path = nueva_carpeta / f'AM_{region}_{fecha_actual}.docx'
    file_path_2 = file_path.as_posix()
    file_path_3 = file_path_2.replace('/', '\\')
    
    doc = word.Documents.Open(file_path_3)
    word.Visible = False
    
    # Actualizar la tabla de contenido (Como es la primera tabla de contenidos estou utilizando "1")
    doc.TablesOfContents(1).Update()
    
    # La seccion que en blanco es la seccion 4
    #section_range = doc.Range(doc.Sections(4).Range.Start, doc.Sections(4).Range.End)
    #section_range.Delete() # eliminar la sección del documento
    
    doc.Close(SaveChanges=True) # guardar y cerrar el documento
    #word.Quit() # cerrar la aplicación de Word
'''
    # Guardar y cerrar el archivo
    #file_path=r'C:\Users\ANALISTAUP29\OneDrive - Ministerio de Educación\MINEDU_2022\GESTION DE LA INFORMACIÓN\UPP\Am Automatizada v2\AM_Automatizada\Am prueba v4.docx'
    #doc.Save(file_path)
    #doc.Close()

    # Cerrar la aplicación Word
    #word.Quit()
    
###########################################################
# Creamos tabla con lista de files para enviar por correo #
###########################################################
    
# Generamos lista de AM.
#lista_AM = glob.glob(os.path.join(proyecto, f"output/AM_corta_region/AM_{fecha_actual}/*"))

#lista_regiones = pd.DataFrame (lista_AM)
#lista_regiones.rename( columns={0:'path'}, inplace=True )
#lista_regiones[['a', 'b', 'c']] = lista_regiones["path"].str.split("AM_", expand = True)
#lista_regiones[['date', 'e']] = lista_regiones["b"].str.split("/", expand = True)
#lista_regiones[['region', 'g']] = lista_regiones["c"].str.split("_", expand = True)
#lista_regiones = lista_regiones[["path", "date","region"]]
#lista_regiones.to_excel(Path(proyecto, "documentacion", "lista_regiones.xlsx"), index = False)
